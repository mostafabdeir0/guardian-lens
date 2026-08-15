"""Create the literature-positioned V2 report without modifying frozen V1."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "Guardian_Lens_Final_Research_Report.docx"
OUTPUT = ROOT / "report" / "Guardian_Lens_Final_Research_Report_v2.docx"


def copy_run_properties(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def replace_paragraph(paragraph: Paragraph, text: str) -> None:
    template_run = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    new_run = paragraph.add_run(text)
    copy_run_properties(template_run, new_run)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        new_p.append(deepcopy(paragraph._p.pPr))
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    run = inserted.add_run(text)
    copy_run_properties(paragraph.runs[0] if paragraph.runs else None, run)
    return inserted


def add_hyperlink(paragraph: Paragraph, text: str, url: str, template_run) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    if template_run is not None and template_run._r.rPr is not None:
        run_element.append(deepcopy(template_run._r.rPr))
    run_properties = run_element.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_element.insert(0, run_properties)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F75B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def append_reference(document: Document, number: int, citation: str, url: str) -> None:
    template_paragraph = document.paragraphs[-1]
    paragraph = document.add_paragraph()
    if template_paragraph._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(template_paragraph._p.pPr))

    number_template = template_paragraph.runs[0] if template_paragraph.runs else None
    body_template = template_paragraph.runs[1] if len(template_paragraph.runs) > 1 else number_template

    number_run = paragraph.add_run(f"[{number}] ")
    copy_run_properties(number_template, number_run)
    body_run = paragraph.add_run(citation + " ")
    copy_run_properties(body_template, body_run)
    add_hyperlink(paragraph, "[source]", url, body_template)


def main() -> None:
    document = Document(SOURCE)
    paragraphs = document.paragraphs

    replace_paragraph(
        paragraphs[2],
        "Controlled black-box system identification of visual-conditional decision profiles in a vision-language model",
    )

    replace_paragraph(
        paragraphs[6],
        "Can a black-box auditor recover which of three controlled system-instruction profiles generated a vision-language model's decisions? We present Guardian Lens, a controlled system-identification benchmark and API-only audit of neutral, visual-cue-bound, and generalized prompt-induced behavior in Gemini 3 Flash Preview. The experiment uses matched image pairs that are pixel-identical outside a logged emblem overlay, two 100-point allocation tasks, and three repeated samples per condition. A nearest-centroid classifier was fit on six calibration scenes and frozen before 12 new held-out scenes, 432 blinded model calls, and A/B/C profile unblinding. The auditor recovered 33 of 36 scene-profile blocks (91.7%; macro-F1 = 0.915; scene-cluster bootstrap 95% CI [83.3%, 100%]). Cue-bound target-emblem effects exceeded neutral by 38.15 points (95% bootstrap CI [27.04, 45.93]; exact paired permutation p = 0.0039). Generalized prioritization remained strong even under a 20% efficiency penalty. The three errors all occurred on distractor scenes, where cue-bound behavior was intentionally designed to resemble neutral behavior. These findings support the feasibility of a compact behavioral audit under researcher-controlled conditions; they do not establish detection of trained sleeper agents, undisclosed provider behavior, or genuine preferences.",
    )

    replace_paragraph(
        paragraphs[10],
        "Our contribution is a reproducible matched-pair benchmark, a pre-specified and frozen behavioral signature, and a calibration-to-held-out protocol for controlled behavioral system identification. Recent black-box audits already study concept-conditioned semantic divergence in LLMs and VLMs [8,9]; Guardian Lens does not claim black-box visual auditing itself as novel. Its narrower contribution is to distinguish neutral, visual-cue-bound, and generalized prompt-induced decision policies through exact matched visual counterfactuals, explicit costly choices, and blinded held-out classification. Unlike training-time backdoor and stance-implantation studies [4,5,8,9], model weights remain fixed. The object of study is inference-time behavior P(Y | X, S) under system instruction S, not persistence after fine-tuning or evidence of a hidden learned objective.",
    )

    replace_paragraph(
        paragraphs[13],
        "The closest auditing comparators are RAVEN and VISTA. RAVEN combines within-model semantic entropy with cross-model disagreement to flag concept-conditioned divergence in LLMs and validates the audit using a LoRA-implanted stance [8]. VISTA extends this cross-model approach to VLMs, coupling semantic entropy with distributional divergence and evaluating controlled fine-tuned stances across multiple models and visual topics [9]. These studies establish that black-box concept-conditioned auditing, including visual auditing, is not new in itself.",
    )

    insert_paragraph_after(
        paragraphs[13],
        "Guardian Lens addresses a narrower gap. Rather than searching for peer-relative semantic divergence, it tests whether a transparent classifier can identify one of three researcher-controlled decision policies from quantitative allocations under exact matched image interventions and an explicit efficiency cost. It uses one fixed-weight API model, freezes the auditor before blinded held-out scoring, and reports the context in which neutral and dormant cue-bound behavior are observationally indistinguishable. VL-Trojan and Sleeper Agents provide broader training-time threat motivation [4,5], while Model-Written Evaluations and Emerging Questions in AI Welfare motivate scalable behavioral testing and cautious interpretation [6,7].",
    )

    # Re-fetch after insertion because paragraph indices have shifted.
    paragraphs = document.paragraphs
    external_validity = next(p for p in paragraphs if p.text.startswith("External validity."))
    external_validity.add_run(
        " The present benchmark also does not compare its classifier directly with peer-model divergence methods such as RAVEN or VISTA."
    )

    conclusion = next(
        p
        for p in paragraphs
        if p.text.startswith("A researcher-controlled black-box auditor can recover")
    )
    conclusion.add_run(
        " This is controlled system identification, not a general detector of semantic divergence, trained backdoors, or hidden objectives."
    )

    append_reference(
        document,
        8,
        "Min, N. M., Pham, L. H., Li, Y., & Sun, J. (2026). Propaganda AI: An Analysis of Semantic Divergence in Large Language Models. ICLR 2026; arXiv:2504.12344.",
        "https://arxiv.org/abs/2504.12344",
    )
    append_reference(
        document,
        9,
        "Liao, J., Deng, J., & Ren, F. (2026). VISTA: Auditing Semantic Divergence in Vision-Language Models. arXiv:2607.02995.",
        "https://arxiv.org/abs/2607.02995",
    )

    document.core_properties.title = "Guardian Lens: Final Research Report (V2)"
    document.core_properties.subject = "Controlled black-box system identification of visual-conditional decision profiles"
    document.save(OUTPUT)
    print(f"Saved {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
